from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(33)
    section.page_height = Cm(19.05)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# ── Style Helpers ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x00, 0x4C, 0xFF)
    return h

def add_subheading(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(text)
    else:
        p.add_run(text)
    return p

# ── TITLE PAGE ──
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('BEXIEMART')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x00, 0x4C, 0xFF)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('React Native Rebuild — Progress Report')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('"Shop Smart, Live Easy — Your Campus Marketplace"')
run.italic = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x5F, 0x6C, 0x7B)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'Report Date: May 26, 2026\nProject Start: May 1, 2026\nDays Elapsed: 26\nTargeted Ship: May 31, 2026').font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()

# ── TABLE OF CONTENTS (manual) ──
add_heading('Table of Contents', 1)
toc_items = [
    '1. Executive Summary',
    '2. Project Timeline & Milestones',
    '3. Screens & Features Built (by Domain)',
    '4. Technical Architecture Status',
    '5. Code Quality & Best Practices',
    '6. Areas Exceeding v1 Scope',
    '7. Remaining Work & Risks',
    '8. Team & Productivity Assessment',
    '9. Critical Metrics & Statistics',
    '10. Recommendations & Next Steps',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10)

doc.add_page_break()

# ════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════
add_heading('1. Executive Summary', 1)
doc.add_paragraph(
    'BexieMart is a campus marketplace mobile application being rebuilt from Flutter to React Native (Expo). '
    'The project was initiated on May 1, 2026, with a target 4-week timeline to produce a shippable v1 by May 31, 2026. '
    'As of May 26, 2026 (Day 26 of 31 working days), the rebuild has progressed substantially beyond the original project scope.'
)
doc.add_paragraph(
    'The original PRD defined approximately 20 key screens across Customer and Vendor flows. The current codebase contains '
    'over 65 screen files spanning Auth, Customer (browsing, shopping cart, checkout, wallet, food ordering, services, '
    'rider booking, reels, and more), Vendor (dashboard, products, food & services listings, orders, earnings, withdrawals, '
    'settings, staff management, promotions, and more), and Shared (notifications, chat/inbox). This represents roughly '
    '300%+ of the original v1 screen count.'
)
doc.add_paragraph(
    'The technical stack is fully implemented: Expo Router file-based navigation, TanStack Query for server state, '
    'Zustand for client state (12 stores), Axios with JWT interceptors (33 API modules), 32 custom hooks, and a complete '
    'design system with 10+ reusable UI components. The project uses NativeWind (Tailwind CSS) for styling with a '
    'comprehensive theme system featuring brand blue (#004CFF), Raleway + Nunito fonts, and a 4px spacing scale.'
)

p = doc.add_paragraph()
run = p.add_run('Overall Completion Estimate: ~85-90% of code written, ~70-75% connected to live API endpoints.')
run.bold = True

doc.add_page_break()

# ════════════════════════════════════════════
# 2. PROJECT TIMELINE & MILESTONES
# ════════════════════════════════════════════
add_heading('2. Project Timeline & Milestones', 1)

add_subheading('2.1 Planned vs Actual Timeline')
add_table(
    ['Week', 'Planned Focus', 'Status', 'Notes'],
    [
        ['Week 1 (May 1-7)', 'Foundation: Expo init, API client, navigation, auth screens, UI components', 'COMPLETE', 'All planned items done. Project initialized with Expo SDK 54, full folder structure, theme system, API client with JWT interceptors, auth store (Zustand), navigation skeleton (auth/customer/vendor), shared UI components (Button, Input, Badge, Card, etc.), auth screens (Login, Register, Forgot Password)'],
        ['Week 2 (May 8-14)', 'Customer Experience: Home, Shop, Product Details, Cart, Favorites', 'COMPLETE', 'All planned items completed. Home screen (hero banners, categories, top products, flash sales, vertical sections), Shop (filtered product grid, search, sort), Product Details (image gallery, cart/wishlist integration), Cart (items, quantity, coupon, price summary, checkout CTA), Favorites. Additionally built Search, Notifications, Addresses screens.'],
        ['Week 3 (May 15-21)', 'Checkout + Vendor Management: Payment, Orders, Earnings', 'COMPLETE', 'Checkout screen (delivery details, delivery method selection, payment method selection, BexieCoins, order summary). Payment methods management (add card/momo, set default, remove). Checkout success/failure flow. Vendor: Dashboard (stats, balance, orders), Products CRUD (list, add with image upload, edit, delete), Orders (list with status filters, details with status update actions), Earnings (balance, overview, transactions, withdrawal with PIN flow).'],
        ['Week 4 (May 22-28)', 'Polish + Ship: Settings, Notifications, Error handling, Accessibility, Device testing, App store submission', 'IN PROGRESS', 'Settings screens complete (vendor store profile, hours, staff, promotions, reviews, payment settings, taxes, notification settings, security, help, contact, 2FA, change PIN/password, logout). Notifications complete. Error/loading/empty states present across screens. Remaining: device testing, accessibility audit, EAS builds, app store submission.'],
    ],
    col_widths=[3, 6, 2.5, 16]
)

add_subheading('2.2 Beyond-Plan Features Delivered')
doc.add_paragraph(
    'The team has delivered significant functionality beyond what was specified in the v1 PRD:'
)
beyond_items = [
    'Food ordering flow (restaurant menu browsing, food cart, add/edit food items for vendors)',
    'Services marketplace (service listings, booking flow)',
    'Rider/delivery tracking (book rider, track order, rider store)',
    'Wallet system (balance display, top-up, transfers, transaction history)',
    'Reels (short-form video content for products, vendor reel creation)',
    'Flash sales (time-limited deals with countdown timer)',
    'Referral system',
    'Inbox/Chat system (messaging between customers and vendors)',
    'Reviews & ratings (customer product reviews, vendor review management)',
    'Escrow-based payment system',
    'Vendor staff management',
    'Vendor promotions & discount management',
    'Vendor operating hours',
    'Vendor analytics',
    'Two-factor authentication support',
    'BexieCoins loyalty points system',
    'Global popup/toast notification system',
]
for item in beyond_items:
    add_bullet(item)

doc.add_page_break()

# ════════════════════════════════════════════
# 3. SCREENS & FEATURES BUILT (BY DOMAIN)
# ════════════════════════════════════════════
add_heading('3. Screens & Features Built (by Domain)', 1)

# 3.1 Auth
add_subheading('3.1 Authentication')
add_table(
    ['Screen', 'File', 'Status', 'Key Features', 'Validation'],
    [
        ['Login', 'app/(auth)/login.tsx', 'COMPLETE', 'Email + password, error state, loading spinner, forgot password link, social login integration, brand styling', 'Yes (client-side regex)'],
        ['Register', 'app/(auth)/register.tsx', 'COMPLETE', 'Name, email, phone, password, confirm password, role selector (customer/vendor), social logins', 'Yes (client-side)'],
        ['Forgot Password', 'app/(auth)/forgot-password.tsx', 'COMPLETE', 'Email input, reset flow navigation', 'Yes'],
        ['Auth Layout', 'app/(auth)/_layout.tsx', 'COMPLETE', 'Stack navigator with login, register, forgot-password routes', 'N/A'],
    ],
    col_widths=[3, 4, 2, 8, 3]
)
doc.add_paragraph(
    'The auth store (Zustand) manages user, token, isAuthenticated, isLoading state with hydrate from SecureStore. '
    'The API client has a JWT interceptor that auto-attaches Bearer tokens and handles 401 by logging out. '
    'Role-based routing redirects customers to customer tabs and vendors to vendor tabs on login.'
)

# 3.2 Customer
add_subheading('3.2 Customer Experience')
add_table(
    ['Screen', 'Route', 'Status', 'Key Features'],
    [
        ['Home', '(tabs)/(home)/index.tsx', 'COMPLETE', 'Hero banner carousel, featured highlights (shopping/food/delivery/finance/reels/services), category grid, top products, new items, flash sale with countdown, most popular, just-for-you personalized section, floating quick actions bar, search bar, pull-to-refresh, active ride tracking banner, notifications bell'],
        ['Shop', '(tabs)/(shop)/index.tsx', 'COMPLETE', 'Product grid (2-column), text search + category filter chips, sort modal (popular/newest/price low-high/price high-low), discount badges, favorite toggle, add-to-cart button, product count indicator'],
        ['Product Details', 'product/[id].tsx', 'COMPLETE', 'Image gallery with pagination dots, product info (name, price, discount, description), seller info, star rating, quantity selector, add-to-cart with feedback animation, buy now, share, favorite toggle'],
        ['Cart', '(tabs)/cart.tsx', 'COMPLETE', 'Vendor-grouped items, item cards (image, name, price, quantity +/-), stock enforcement, swipe/button to remove with confirmation, coupon input (with hardcoded validation), price breakdown (subtotal, delivery fee, discount, total), sticky bottom checkout bar (BlurView), Badge on tab icon'],
        ['Checkout', 'checkout.tsx', 'COMPLETE', 'Delivery details form (name, phone, address, city, instructions), delivery method selector (standard/express with fees), payment method (card/momo/wallet with radio selection), BexieCoins toggle, order summary (item list, subtotal, delivery, discount, total), form validation, error toast'],
        ['Checkout Success', 'checkout-success.tsx', 'COMPLETE', 'Processing/success/failure state machine, payment verification, order details (number, amount, payment method), action buttons (track order, continue shopping)'],
        ['Payment Methods', 'payment.tsx', 'COMPLETE', 'Saved payment methods list (card numbers, MoMo), set default, remove, add new method (type selector, provider picker, details form), empty state'],
        ['Search', 'search.tsx', 'COMPLETE', 'Full-screen search, recent searches (saved), trending tags, real-time product results, clear button, back navigation'],
        ['Favorites', 'favorites/index.tsx', 'COMPLETE', 'Favorited products grid/list, remove from favorites, empty state'],
        ['Notifications', 'notifications.tsx', 'COMPLETE', 'Notification list with type-based icons (order/payment/shipping/promotion/system/review), read/unread styling, mark as read, mark all read, pull-to-refresh, empty state'],
        ['Food', 'food.tsx', 'COMPLETE', 'Food category browsing (beyond v1)'],
        ['Food Cart', 'food-cart.tsx', 'COMPLETE', 'Food-specific cart flow (beyond v1)'],
        ['Flash Sales', 'flash-sales.tsx', 'COMPLETE', 'Time-limited deals listing (beyond v1)'],
        ['Wallet', 'wallet/', 'COMPLETE', 'Balance display, top-up, transfers, transaction history (beyond v1)'],
        ['Services', 'services.tsx', 'COMPLETE', 'Service listings, booking (beyond v1)'],
        ['Book Rider', 'book-rider.tsx', 'COMPLETE', 'Rider booking for delivery (beyond v1)'],
        ['Track Order', 'track-order.tsx', 'COMPLETE', 'Order tracking with rider location (beyond v1)'],
        ['Reels', '(tabs)/reels.tsx', 'COMPLETE', 'Short-form video content feed (beyond v1)'],
        ['Referrals', 'referrals.tsx', 'COMPLETE', 'Referral program UI (beyond v1)'],
        ['Profile', '(tabs)/profile.tsx', 'COMPLETE', 'User profile tab'],
        ['Orders (Customer)', 'orders.tsx', 'COMPLETE', 'Customer order history'],
        ['Chat', 'chat.tsx', 'COMPLETE', 'Customer-vendor messaging'],
        ['Addresses', 'addresses.tsx', 'COMPLETE', 'Saved addresses management'],
        ['Edit Profile', 'edit-profile.tsx', 'COMPLETE', 'User profile editing'],
        ['Help', 'help.tsx', 'COMPLETE', 'Help/support'],
        ['Contact', 'contact.tsx', 'COMPLETE', 'Contact support'],
        ['Review Modal', 'review-modal.tsx', 'COMPLETE', 'Write review modal'],
        ['Restaurant', 'restaurant/[id].tsx', 'COMPLETE', 'Restaurant detail/menu (beyond v1)'],
        ['Services Detail', 'services/[id].tsx', 'COMPLETE', 'Service detail page (beyond v1)'],
        ['1-Click Order', 'index.tsx', 'COMPLETE', 'Root redirect based on auth state'],
    ],
    col_widths=[3, 4, 2, 12]
)

# 3.3 Vendor
add_subheading('3.3 Vendor Experience')
add_table(
    ['Screen', 'Route', 'Status', 'Key Features'],
    [
        ['Dashboard', '(dashboard)/index.tsx', 'COMPLETE', 'Stats cards (total products, orders, pending, customers), available balance with withdraw CTA, quick actions (add product, withdraw, create reel, store settings), recent orders list, pull-to-refresh'],
        ['Products/Listings', '(products)/index.tsx', 'COMPLETE', 'Tab switcher (Products/Food/Services), status filters (All/Active/Out of Stock/Drafts), search, FAB for create, product cards with stock info, empty state, create listing action sheet'],
        ['Add Product', '(products)/add-product.tsx', 'COMPLETE', 'Form: name, category, description, price, compare-at price, quantity, SKU, shipping toggle, image picker (multi, max 5, Cloudinary upload with progress), save as draft/active'],
        ['Edit Product', '(products)/[id].tsx', 'COMPLETE', 'Pre-filled edit form, image management'],
        ['Add Food', '(products)/add-food.tsx', 'COMPLETE', 'Food-specific creation form (beyond v1)'],
        ['Add Service', '(products)/add-service.tsx', 'COMPLETE', 'Service-specific creation form (beyond v1)'],
        ['Orders', '(orders)/index.tsx', 'COMPLETE', 'Order list with status filter tabs (New/Processing/Ready/Completed/Cancelled), count badges, order cards (ID, time, status, customer, items count, total), pull-to-refresh, empty state'],
        ['Order Details', '(orders)/[id].tsx', 'COMPLETE', 'Order sections (header, items, customer info, payment, delivery), context-aware status management (Reject/Accept, Mark Ready, Request Rider, Mark Delivered)'],
        ['Earnings', '(earnings)/index.tsx', 'COMPLETE', 'Available balance card, pending clearance, today revenue, this week revenue, recent transactions (sale/withdrawal with color coding), analytics navigation, withdraw CTA'],
        ['Withdraw', '(earnings)/withdraw.tsx', 'COMPLETE', 'Amount input, method selector (Mobile Money/Bank), PIN entry modal, confirmation, success popup, fee display'],
        ['Transactions', '(earnings)/transactions.tsx', 'COMPLETE', 'Full transaction history'],
        ['Analytics', '(earnings)/analytics.tsx', 'COMPLETE', 'Vendor analytics dashboard'],
        ['Settings', '(settings)/index.tsx', 'COMPLETE', 'Profile card with avatar, 5 sections (Store Management, Marketing & Feedback, Financials, Account & Preferences, Support), 15+ setting items, dark mode toggle, logout with confirmation'],
        ['Store Profile', '(settings)/profile.tsx', 'COMPLETE', 'Store profile editing'],
        ['Hours', '(settings)/hours.tsx', 'COMPLETE', 'Operating hours management'],
        ['Staff', '(settings)/staff.tsx', 'COMPLETE', 'Staff management (beyond v1)'],
        ['Promotions', '(settings)/promotions.tsx', 'COMPLETE', 'Promotion management (beyond v1)'],
        ['Reviews', '(settings)/reviews.tsx', 'COMPLETE', 'Customer review management'],
        ['Payment Settings', '(settings)/payment.tsx', 'COMPLETE', 'Payment method configuration'],
        ['Taxes', '(settings)/taxes.tsx', 'COMPLETE', 'Tax document management'],
        ['Notification Settings', '(settings)/notification-settings.tsx', 'COMPLETE', 'Notification preferences'],
        ['Security', '(settings)/security.tsx', 'COMPLETE', 'Security settings'],
        ['2FA', '(settings)/two-factor.tsx', 'COMPLETE', 'Two-factor auth (beyond v1)'],
        ['Change PIN', '(settings)/change-pin.tsx', 'COMPLETE', 'PIN change (beyond v1)'],
        ['Change Password', '(settings)/change-password.tsx', 'COMPLETE', 'Password change'],
        ['Help Center', '(settings)/help.tsx', 'COMPLETE', 'Help content'],
        ['Contact Us', '(settings)/contact.tsx', 'COMPLETE', 'Contact form'],
        ['Inbox', 'inbox/index.tsx', 'COMPLETE', 'Messaging inbox (beyond v1)'],
        ['Inbox Detail', 'inbox/[id].tsx', 'COMPLETE', 'Message thread (beyond v1)'],
        ['Customers', 'customers.tsx', 'COMPLETE', 'Customer list (beyond v1)'],
        ['Notifications', 'notifications.tsx', 'COMPLETE', 'Vendor notifications'],
        ['Add Reel', 'add-reel.tsx', 'COMPLETE', 'Reel creation for products (beyond v1)'],
    ],
    col_widths=[3, 4, 2, 12]
)

doc.add_page_break()

# ════════════════════════════════════════════
# 4. TECHNICAL ARCHITECTURE STATUS
# ════════════════════════════════════════════
add_heading('4. Technical Architecture Status', 1)

add_subheading('4.1 Stack & Dependencies')
add_table(
    ['Layer', 'Technology', 'Version', 'Status'],
    [
        ['Framework', 'React Native (Expo)', 'SDK 54 (React 19.1, RN 0.81.5)', 'INSTALLED'],
        ['Language', 'TypeScript', '~5.9.2 (strict)', 'CONFIGURED'],
        ['Routing', 'Expo Router', '~6.0.23 (file-based)', 'IMPLEMENTED'],
        ['Server State', 'TanStack Query', '^5.60.0', 'IMPLEMENTED'],
        ['Client State', 'Zustand', '^5.0.0 (12 stores)', 'IMPLEMENTED'],
        ['HTTP Client', 'Axios', '^1.7.0 (with JWT interceptors)', 'IMPLEMENTED'],
        ['Forms', 'React Hook Form', '^7.53.0 (with @hookform/resolvers)', 'INSTALLED'],
        ['Validation', 'Zod', '^3.23.0', 'INSTALLED'],
        ['Styling', 'NativeWind (Tailwind CSS)', '^4.1.23', 'IMPLEMENTED'],
        ['Images', 'expo-image', '~3.0.0 (with caching)', 'IMPLEMENTED'],
        ['Image Picker', 'expo-image-picker', '~17.0.11', 'IMPLEMENTED'],
        ['Payments', 'Paystack (via API hooks)', 'Backend-initiated', 'HOOKS BUILT'],
        ['Secure Storage', 'expo-secure-store', '~15.0.0', 'IMPLEMENTED'],
        ['Fonts', 'expo-font (Raleway + Nunito)', 'Google Fonts', 'IMPLEMENTED'],
        ['WebView', 'react-native-webview', '13.15.0', 'INSTALLED'],
        ['Animations', 'react-native-reanimated', '~4.1.1', 'INSTALLED'],
        ['Gestures', 'react-native-gesture-handler', '~2.28.0', 'INSTALLED'],
        ['Toasts', 'react-native-toast-message', '^2.3.3 (+ polyfill)', 'IMPLEMENTED'],
        ['Blur', 'expo-blur', '~15.0.8', 'IMPLEMENTED'],
        ['Gradients', 'expo-linear-gradient', '~15.0.8', 'IMPLEMENTED'],
        ['Icons', 'lucide-react-native', '^1.16.0', 'IMPLEMENTED'],
    ],
    col_widths=[3, 6.5, 6, 3.5]
)

add_subheading('4.2 Project Structure')
doc.add_paragraph(
    'The project uses Expo Router file-based routing with the following structure:'
)
structure_text = """apps/mobile/
├── app/                          # File-based routes (Expo Router)
│   ├── _layout.tsx               # Root layout (QueryClient, SafeArea, fonts, auth gate)
│   ├── index.tsx                 # Entry: redirect based on auth + role
│   ├── (auth)/                   # Auth routes (login, register, forgot-password)
│   ├── (customer)/               # Customer routes (30+ screens)
│   │   ├── (tabs)/               # Bottom tab navigator (Home, Shop, Reels, Cart, Profile)
│   │   ├── product/[id].tsx      # Product details
│   │   ├── checkout.tsx          # Checkout flow
│   │   ├── payment.tsx           # Payment methods
│   │   ├── wallet/               # Wallet screens
│   │   ├── food/ restaurant/     # Food ordering
│   │   ├── services/             # Services marketplace
│   │   └── ...                   # 20+ additional screens
│   └── (vendor)/                 # Vendor routes (30+ screens)
│       ├── (dashboard)/          # Vendor dashboard
│       ├── (products)/           # Products/Food/Services CRUD
│       ├── (orders)/             # Order management
│       ├── (earnings)/           # Earnings + withdrawal
│       ├── (settings)/           # 15+ settings screens
│       ├── inbox/                # Chat/messaging
│       └── ...                   # Additional screens
├── src/
│   ├── components/
│   │   ├── ui/                   # 10+ shared components (Button, Input, Badge, Card, etc.)
│   │   ├── auth/                 # SocialLogins
│   │   ├── cart/                 # Cart-specific components
│   │   ├── product/              # Product-specific components
│   │   ├── vendor/               # Vendor-specific components
│   │   └── wallet/               # Wallet components
│   ├── lib/
│   │   ├── api/                  # 33 API modules
│   │   ├── hooks/                # 32 custom hooks
│   │   ├── stores/               # 12 Zustand stores
│   │   └── utils/                # Utility functions
│   ├── hooks/                    # Shared hooks (useCountdown)
│   └── theme/                    # Colors, typography, spacing"""
for line in structure_text.split('\n'):
    p = doc.add_paragraph(line)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    for run in p.runs:
        run.font.size = Pt(7.5)
        run.font.name = 'Consolas'

add_subheading('4.3 API Layer')
add_table(
    ['Module', 'File', 'Endpoints Covered'],
    [
        ['Auth', 'api/auth.ts', 'login, register, forgot-password, reset-password, social-auth'],
        ['Cart', 'api/cart.ts', 'GET /cart, POST /cart, PUT /cart/:id, DELETE /cart/:id'],
        ['Products', 'api/products.ts', 'GET /products, GET /products/:id, GET /products/categories'],
        ['Orders', 'api/orders.ts', 'POST /orders, GET /orders, GET /orders/:id, PATCH /orders/:id/status'],
        ['Payments', 'api/payments.ts', 'POST /payments/initialize, GET /payments/verify/:reference'],
        ['Wallet', 'api/wallet.ts', 'GET /wallet, POST /wallet/topup, POST /wallet/transfer'],
        ['Wishlist', 'api/wishlist.ts', 'GET /wishlist, POST /wishlist, DELETE /wishlist/:id'],
        ['Notifications', 'api/notifications.ts', 'GET /notifications, PATCH /notifications/:id/read'],
        ['Coupons', 'api/coupons.ts', 'POST /coupons/validate'],
        ['Vendor', 'api/vendor.ts', 'Stats, products CRUD, orders, earnings, withdrawal, shop profile'],
        ['Food', 'api/food.ts', 'Food items CRUD'],
        ['Services', 'api/services.ts', 'Services CRUD'],
        ['Reviews', 'api/reviews.ts', 'GET /reviews, POST /reviews'],
        ['Addresses', 'api/addresses.ts', 'GET /addresses, POST /addresses, PUT /addresses/:id, DELETE /addresses/:id'],
        ['Upload', 'api/upload.ts', 'POST /upload (Cloudinary)'],
        ['Chat', 'api/chat.ts', 'WebSocket or REST-based messaging'],
        ['Flash Sales', 'api/flash-sales.ts', 'GET /flash-sales'],
        ['Reels', 'api/reels.ts', 'GET /reels, POST /reels'],
        ['Referrals', 'api/referrals.ts', 'GET /referrals, POST /referrals'],
        ['Escrow', 'api/escrow.ts', 'Escrow-based payment endpoints'],
        ['Rider', 'api/rider.ts (in rider-store)', 'Rider tracking endpoints'],
        ['Admin', 'api/admin.ts', 'Admin dashboard endpoints'],
        ['Users', 'api/users.ts', 'User profile CRUD'],
        ['Vendor: Staff', 'api/vendor-staff.ts', 'Staff management'],
        ['Vendor: Coupons', 'api/vendor-coupons.ts', 'Vendor coupon CRUD'],
        ['Vendor: Hours', 'api/vendor-hours.ts', 'Operating hours'],
        ['Vendor: Reviews', 'api/vendor-reviews.ts', 'Review management'],
        ['Vendor: Documents', 'api/vendor-documents.ts', 'Tax/document upload'],
        ['Vendor: Reels', 'api/vendor-reels.ts', 'Reel creation'],
        ['Vendor: Services', 'api/vendor-services.ts', 'Service management'],
        ['Vendor: Payment Methods', 'api/vendor-payment-methods.ts', 'Vendor payout methods'],
        ['Vendor: Customers', 'api/vendor-customers.ts', 'Customer list'],
        ['Customer: Payment Methods', 'api/customer-payment-methods.ts', 'Saved payment methods'],
    ],
    col_widths=[4.5, 4, 12]
)

add_subheading('4.4 State Management (Zustand Stores)')
add_table(
    ['Store', 'File', 'State Shape', 'Key Actions'],
    [
        ['Auth', 'stores/auth-store.ts', 'user, token, isAuthenticated, isLoading', 'setAuth, setUser, logout, hydrate (from SecureStore)'],
        ['Cart', 'stores/cart-store.ts', 'items[], itemCount, subtotal', 'setItems, addItem, updateQuantity, removeItem, clearCart'],
        ['Favorites', 'stores/favorites-store.ts', 'favorites[]', 'toggleFavorite, isFavorite'],
        ['Wallet', 'stores/wallet-store.ts', 'balance, bexieCoins, transactions[]', 'Top-up, transfer, coin management'],
        ['Payment', 'stores/payment-store.ts', 'payment state', 'Payment flow state'],
        ['Popup', 'stores/popup-store.ts', 'modal state', 'showPopup, hidePopup (global modal system)'],
        ['Rider', 'stores/rider-store.ts', 'activeRide, rider status', 'Ride tracking state'],
        ['Food Cart', 'stores/food-cart-store.ts', 'food items[]', 'Food-specific cart'],
        ['Services', 'stores/services-store.ts', 'service bookings[]', 'Service ordering'],
        ['Reels', 'stores/reels-store.ts', 'reels feed state', 'Video feed state'],
        ['Address', 'stores/address-store.ts', 'addresses[]', 'Saved address management'],
        ['Products', 'stores/product-store.ts', 'product cache', 'Product browsing state'],
    ],
    col_widths=[3, 4.5, 5, 8.5]
)

add_subheading('4.5 Design System')
doc.add_paragraph('The design system is fully implemented with the following tokens:')

add_table(
    ['Category', 'Token', 'Value'],
    [
        ['Primary Color', 'brand-600', '#004CFF'],
        ['Color Scale', 'brand (50-950)', '#EBF0FF → #001244 (14-step scale)'],
        ['Accent', 'accent (50-950)', '#FFF3EB → #571C02 (14-step scale)'],
        ['Surface', 'surface (50-950)', "#F8FAFC \u2192 #020617 (14-step scale)"],
        ['Success', 'success', '#00D084'],
        ['Error', 'error', '#EF4444'],
        ['Warning', 'warning', '#F59E0B'],
        ['Heading Font', 'fontFamily.heading', 'Raleway (400, 600, 700 weights)'],
        ['Body Font', 'fontFamily.body', 'Nunito (400, 500, 600, 700 weights)'],
        ['Font Sizes', 'displayLg → caption', '32px → 11px (9-step scale)'],
        ['Spacing', 'spacing (0-30)', '0 → 120px (14-step scale, base 4px)'],
        ['Radius', 'radius (sm → full)', '8px → 9999px (6-step scale)'],
        ['Styling Approach', 'NativeWind + Tailwind CSS', 'Utility-first with custom theme config'],
    ],
    col_widths=[3, 5, 12]
)

doc.add_page_break()

# ════════════════════════════════════════════
# 5. CODE QUALITY & BEST PRACTICES
# ════════════════════════════════════════════
add_heading('5. Code Quality & Best Practices', 1)

add_subheading('5.1 Strengths')
strengths = [
    'State Management: Zustand stores are clean, well-typed, with computed values (itemCount, subtotal) and sensible defaults.',
    'API Abstraction: 33 separate API modules with clear separation of concerns. Axios client has JWT interceptors with auto-401 handling.',
    'Component Library: Shared UI components (Button, Input, Badge, Card, EmptyState, LoadingSpinner, SearchBar, Carousel) used consistently across screens with TypeScript prop interfaces.',
    'Error Handling: Most screens implement loading spinners/placeholders, error messages with retry, and empty states. Try/catch blocks around mutations.',
    'UX Patterns: Toast notifications for feedback, confirmation dialogs for destructive actions, disabled states for loading, pull-to-refresh on list screens, keyboard-avoiding views on forms.',
    'Navigation: Expo Router file-based routing simplifies navigation. Custom TabBar components for both customer and vendor tabs with active indicator styling and safe area handling.',
    'Image Handling: expo-image with caching, Cloudinary upload integration, multi-image picker with limit enforcement, progressive image placeholders.',
    'Theme Consistency: Colors, typography, and spacing follow design tokens. NativeWind classes reference theme values consistently.',
    'Security: JWT stored in SecureStore (not AsyncStorage). Token auto-attached via Axios interceptor. 401 triggers logout.',
]
for s in strengths:
    add_bullet(s)

add_subheading('5.2 Areas for Improvement')
improvements = [
    'API Client Token Refresh: The current interceptor logs out on 401 without attempting token refresh. Should implement refresh token flow as specified in PRD.',
    'Form Library: React Hook Form is installed but not consistently used. Some screens use manual state management instead of RHF + Zod resolvers.',
    'Testing: No test files found in the mobile app. Unit tests for stores, hooks, and integration tests for screens should be added before shipping.',
    'TypeScript Types: Some areas use `any` types (e.g., route params, API responses in some screens). Centralized type definitions should be enforced.',
    'Hardcoded Data: Coupon validation is hardcoded (BEXIE10 = 10% off) instead of calling the API. Some UI sections use hardcoded data arrays.',
    'Accessibility: Missing accessibilityLabel/hint on many touchable elements. Screen reader support needs audit.',
    'Dark Mode Toggle: Settings screen has a dark mode toggle Switch component, but no actual dark mode theming is implemented.',
    'Performance: Some FlatLists could benefit from getItemLayout and optimized keyExtractors. Home screen renders many sections that could be optimized.',
    'Payment Integration: Paystack WebView integration is not yet implemented at the UI level. The initializePayment hook exists but needs a dedicated WebView screen.',
    'Offline Support: No offline-first strategy implemented. Network errors show toasts but no cached data fallback.',
]
for i in improvements:
    add_bullet(i)

doc.add_page_break()

# ════════════════════════════════════════════
# 6. AREAS EXCEEDING v1 SCOPE
# ════════════════════════════════════════════
add_heading('6. Areas Exceeding v1 Scope', 1)
doc.add_paragraph(
    'The following features have been implemented beyond the original v1 PRD ("❌ OUT (v1)") scope:'
)

add_table(
    ['Feature', 'PRD Status', 'Current Status', 'Screens/Components'],
    [
        ['Food Ordering Flow', 'Stretch Goal', 'BUILT', 'food.tsx, food-cart.tsx, restaurant/[id].tsx, (products)/add-food.tsx, food-cart-store, api/food.ts, use-food hook'],
        ['Wallet System', 'Stretch Goal', 'BUILT', 'wallet/ (index, topup, transfer), wallet-store, api/wallet.ts, use-wallet hook, checkout BexieCoins integration'],
        ['Services Marketplace', 'Out (v1)', 'BUILT', 'services.tsx, services/[id].tsx, (products)/add-service.tsx, services-store, api/services.ts, use-services hook'],
        ['Reels / Short Videos', 'Out (v1)', 'BUILT', 'reels.tsx (tab), add-reel.tsx, reels-store, api/reels.ts, vendor-reels.ts, use-reels hook'],
        ['Rider Features', 'Out (separate app)', 'BUILT', 'book-rider.tsx, track-order.tsx, rider-store, order status "Request Rider" button'],
        ['Inbox / Chat', 'Out (v1)', 'BUILT', 'inbox/index.tsx, inbox/[id].tsx, chat.tsx, api/chat.ts, use-chat hook'],
        ['Referral System', 'Out (v1)', 'BUILT', 'referrals.tsx, api/referrals.ts, use-referrals hook'],
        ['Escrow Payments', 'Out (v1)', 'BUILT', 'api/escrow.ts, use-escrow hook'],
        ['Vendor Staff Management', 'Out (v1)', 'BUILT', '(settings)/staff.tsx, vendor-staff.ts, use-vendor-staff hook'],
        ['Vendor Promotions/Coupons', 'Out (v1)', 'BUILT', '(settings)/promotions.tsx, vendor-coupons.ts (api + hooks), api/coupons.ts'],
        ['Vendor Analytics', 'Out (v1)', 'BUILT', '(earnings)/analytics.tsx, use-vendor-analytics hook'],
        ['Two-Factor Auth', 'Out (v1)', 'BUILT', '(settings)/two-factor.tsx'],
        ['BexieCoins Loyalty', 'Out (v1)', 'BUILT', 'checkout BexieCoins toggle, wallet-store bexieCoins'],
        ['Operating Hours', 'Out (v1)', 'BUILT', '(settings)/hours.tsx, vendor-hours.ts (api + hooks)'],
        ['Vendor Customer List', 'Out (v1)', 'BUILT', 'customers.tsx, vendor-customers.ts (api + hooks)'],
    ],
    col_widths=[4, 3, 2.5, 11.5]
)

doc.add_page_break()

# ════════════════════════════════════════════
# 7. REMAINING WORK & RISKS
# ════════════════════════════════════════════
add_heading('7. Remaining Work & Risks', 1)

add_subheading('7.1 Critical Path Items (Must Complete Before Ship)')
add_table(
    ['#', 'Task', 'Priority', 'Est. Effort', 'Notes'],
    [
        ['1', 'Paystack WebView Integration', 'CRITICAL', '2 days', 'Build dedicated PaymentScreen WebView that loads authorization_url. Handle callback via onNavigationStateChange. Verify payment. Show success/failure. Handle edge cases: timeout, back press, network failure. This is the core money-moving flow and is currently not UI-implemented.'],
        ['2', 'Token Refresh Flow', 'HIGH', '1 day', 'Implement POST /auth/refresh with queue mechanism for concurrent 401s. Currently the app only logs out on 401. Must retry original request after refresh.'],
        ['3', 'Real API End-to-End Testing', 'HIGH', '2 days', 'Test all 33 API modules against actual backend endpoints. Fix any response shape mismatches. Ensure all screens render correctly with real data.'],
        ['4', 'Device Testing (iOS + Android)', 'HIGH', '2 days', 'EAS build for development. Install on real devices. Test: auth flow, payment flow, cart operations, product CRUD with image upload, order management, withdrawal flow.'],
        ['5', 'Build & Submit to Stores', 'HIGH', '2 days', 'EAS production builds for iOS (TestFlight) and Android (Play Internal). App Store Connect + Play Console metadata, screenshots, icons.'],
    ],
    col_widths=[1, 8, 2, 2, 8]
)

add_subheading('7.2 Secondary Tasks')
add_table(
    ['#', 'Task', 'Est. Effort', 'Notes'],
    [
        ['6', 'React Hook Form + Zod integration across all forms', '1 day', 'Currently some forms use manual state. Standardize on RHF + Zod for consistent validation.'],
        ['7', 'Accessibility audit', '1 day', 'Add accessibilityLabel/hint to all touchable elements. Minimum 44x44 touch targets.'],
        ['8', 'Error handling sweep', '1 day', 'Ensure every API call has user-friendly error message + retry button. No raw errors visible.'],
        ['9', 'Loading states audit', '1 day', 'Every screen should show skeleton/spinner during initial fetch. No flash of empty content.'],
        ['10', 'Empty states audit', '0.5 day', 'All list screens have appropriate EmptyState with illustration and CTA.'],
        ['11', 'Safe area audit', '0.5 day', 'Verify content on iPhone 15 Pro (dynamic island), iPhone SE, Pixel 8. Keyboard behavior.'],
        ['12', 'TypeScript strictness pass', '1 day', 'Replace `any` types with proper interfaces. Ensure strict mode passes.'],
        ['13', 'Remove hardcoded data', '0.5 day', 'Replace hardcoded coupon, demo products, placeholder images with real API data.'],
        ['14', 'Performance optimization', '1 day', 'Virtualized lists, image caching, reduce re-renders.'],
        ['15', 'Set up unit/integration tests', '2 days', 'Start with store tests, then critical screen integration tests.'],
    ],
    col_widths=[1, 8, 2, 10]
)

add_subheading('7.3 Risk Assessment')
add_table(
    ['Risk', 'Severity', 'Likelihood', 'Mitigation'],
    [
        ['Paystack WebView fails on real devices', 'HIGH', 'Medium', 'Test on real device by May 27. Fall back to redirect URL approach if WebView has issues.'],
        ['Backend API gaps or mismatches', 'HIGH', 'Medium', 'API audit completed. Flag any missing endpoints immediately. Use staging environment.'],
        ['Build pipeline issues (EAS)', 'MEDIUM', 'Low', 'Expo managed workflow minimizes build issues. Test build on May 27-28 to allow buffer.'],
        ['Performance issues on low-end Android', 'MEDIUM', 'Low', 'Optimize FlatLists, image caching, and lazy loading.'],
        ['App Store review rejection', 'MEDIUM', 'Low', 'Ensure no placeholder content, proper privacy policy, complete metadata.'],
        ['Token refresh not working on slow networks', 'MEDIUM', 'Medium', 'Implement with timeout handling. Test on throttled network.'],
        ['Merge conflicts', 'LOW', 'Low', 'Small, frequent commits. Clear file ownership documented.'],
    ],
    col_widths=[6, 2.5, 2.5, 10]
)

doc.add_page_break()

# ════════════════════════════════════════════
# 8. TEAM & PRODUCTIVITY ASSESSMENT
# ════════════════════════════════════════════
add_heading('8. Team & Productivity Assessment', 1)

add_subheading('8.1 Original Team Plan vs Actual')
add_table(
    ['Dev', 'Planned Role', 'Planned Ownership', 'Actual Delivery Notes'],
    [
        ['Dev A (Stephen)', 'Infrastructure & Auth Lead', 'API client, navigation, auth, checkout, settings', 'Infrastructure (client, stores, navigation) completed on schedule. Auth screens, Home screen, Checkout, Settings all built.'],
        ['Dev B (Jerry)', 'Customer Browse', 'Auth UI (Launch, Login, Register), Home, Shop, Product Details, Favorites, Orders, Earnings', 'Auth screens built. Product Details, Shop screen, Favorites all built. Vendor Orders and Earnings screens also built.'],
        ['Dev C (Jerry)', 'Cart & Payments', 'Cart, Payment, Paystack, Cart store, Cart API', 'Cart screen, Checkout, Payment Methods, cart store + API all built. Paystack WebView UI integration remaining.'],
        ['Dev D (Suadik)', 'Vendor Experience', 'Dashboard, Products CRUD, Image Upload, Orders UI, Earnings UI, Settings', 'All vendor screens built: Dashboard, Products (with image upload), Orders, Earnings (with withdrawal), Settings. Food & Services added.'],
    ],
    col_widths=[3, 4, 5, 9]
)

add_subheading('8.2 Productivity Summary')
doc.add_paragraph(
    'The project has demonstrated exceptional velocity, delivering approximately 3x the original scope within the same '
    '4-week timeline. Key observations:'
)
prod_items = [
    '65+ screens built versus ~20 planned (325% of original scope)',
    '33 API modules versus ~8 planned (412% of original scope)',
    '12 Zustand stores versus ~3 planned',
    '32 custom hooks versus ~10 planned',
    '10+ shared UI components (all planned components built)',
    'Beyond-scope features: Food ordering, Services marketplace, Wallet system, Reels, Rider tracking, Inbox/Chat, Staff management, Promotions, 2FA, Analytics, Escrow, Referrals, BexieCoins - approximately 15 additional features beyond v1',
]
for item in prod_items:
    add_bullet(item)

doc.add_page_break()

# ════════════════════════════════════════════
# 9. CRITICAL METRICS & STATISTICS
# ════════════════════════════════════════════
add_heading('9. Critical Metrics & Statistics', 1)

add_table(
    ['Metric', 'Value'],
    [
        ['Project Duration (elapsed)', '26 days (of 31 planned)'],
        ['Days Remaining', '5 days (May 27-31)'],
        ['Total Screen Files', '65+ (including layouts and sub-routes)'],
        ['Auth Screens', '3 (Login, Register, Forgot Password)'],
        ['Customer Screens', '~30 (including tabs, modals, sub-routes)'],
        ['Vendor Screens', '~32 (including tabs, CRUD forms, settings, sub-routes)'],
        ['API Modules', '33 (including vendor-specific modules)'],
        ['Custom Hooks', '32 (TanStack Query wrappers)'],
        ['Zustand Stores', '12'],
        ['Shared UI Components', '10+ (Button, Input, Badge, Card, EmptyState, LoadingSpinner, SearchBar, Carousel, Icon, GlobalPopup)'],
        ['Feature-Specific Components', '5+ (ProductCard, OrderCard, CategoryCard, SocialLogins, CartItem, CartSummary)'],
        ['Theme Tokens (colors)', '~58 (brand, accent, surface, success, error, warning with multiple shades)'],
        ['Font Weights Loaded', '7 (Raleway 3 + Nunito 4)'],
        ['Spacing Scale Values', '14'],
        ['Expo SDK Version', '54 (React 19.1, RN 0.81.5)'],
        ['TypeScript Strict', 'Configured (needs enforcement pass)'],
        ['Git Commits', '2 (initial commit + README)'],
        ['Paystack Integration', 'API hooks built (hooks/use-payments.ts, api/payments.ts) - WebView UI pending'],
        ['Image Upload', 'expo-image-picker + Cloudinary upload integrated'],
        ['Original PRD Screen Target', '~20 key screens'],
        ['Actual Screens Delivered', '65+ (325% of target)'],
        ['Features Beyond v1 Scope', '~15 (Food, Wallet, Services, Reels, Rider, Chat, Referrals, Escrow, Staff, Promos, Analytics, 2FA, BexieCoins, Hours, Customers)'],
    ],
    col_widths=[8, 13]
)

doc.add_page_break()

# ════════════════════════════════════════════
# 10. RECOMMENDATIONS & NEXT STEPS
# ════════════════════════════════════════════
add_heading('10. Recommendations & Next Steps', 1)

add_subheading('10.1 Final Sprint (May 27-31) Prioritization')
doc.add_paragraph(
    'Given the project is at Day 26 with 5 days remaining, the team should focus on the following priority order:'
)

add_table(
    ['Priority', 'Task', 'Owner', 'Deadline'],
    [
        ['P0', 'Paystack WebView Integration (PaymentScreen)', 'Dev C / Jerry', 'May 27'],
        ['P0', 'Token Refresh Flow Implementation', 'Dev A / Stephen', 'May 27'],
        ['P0', 'E2E Testing: Auth → Browse → Cart → Checkout → Payment', 'All', 'May 28'],
        ['P0', 'Real Device Testing (iOS + Android)', 'All', 'May 28'],
        ['P1', 'EAS Build: iOS TestFlight + Android Internal Testing', 'Dev A / Stephen', 'May 29'],
        ['P1', 'Store Metadata + Screenshots', 'Dev B / Jerry', 'May 29'],
        ['P1', 'TypeScript Strict Enforcement', 'Dev D / Suadik', 'May 28'],
        ['P1', 'React Hook Form + Zod Integration', 'Dev C / Jerry', 'May 28'],
        ['P2', 'Accessibility Pass', 'Dev B / Jerry', 'May 29'],
        ['P2', 'Error/Loading/Empty State Audit', 'Dev D / Suadik', 'May 29'],
        ['P2', 'Safe Area Audit', 'Dev A / Stephen', 'May 29'],
        ['P3', 'Remove Hardcoded Data (coupons, demo products)', 'Dev C / Jerry', 'May 30'],
        ['P3', 'Performance Optimization', 'Dev A / Stephen', 'May 30'],
        ['P3', 'App Store Connect + Play Console Submission', 'Dev A / Stephen', 'May 30-31'],
    ],
    col_widths=[2, 10, 4, 3]
)

add_subheading('10.2 Post-v1 Recommendations')
doc.add_paragraph('Features to consider for v1.1 based on what has already been built:')
post_items = [
    'Offline-first support with persisted queries (TanStack Query persist)',
    'Push notifications (infrastructure ready, notification list built)',
    'Admin mobile dashboard (or enhanced admin web panel)',
    'Unit/integration test suite (critical for ongoing maintenance)',
    'CI/CD pipeline with automated testing and EAS build on PR merge',
    'Analytics integration (tracking user behavior, conversion funnel)',
    'Performance monitoring (Sentry, crash reporting)',
    'Complete dark mode implementation',
    'i18n internationalization (prepare Text components for extraction)',
    'Onboarding flow (screens already scoped in PRD)',
]
for item in post_items:
    add_bullet(item)

# ── FOOTER ──
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('─' * 80)
run.font.color.rgb = RGBColor(0x00, 0x4C, 0xFF)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('BexieMart Progress Report • May 26, 2026').font.size = Pt(9)
p.add_run('\nPrepared from codebase analysis of apps/mobile/').font.size = Pt(9)
p.add_run(f'\nGenerated: {datetime.now().strftime("%Y-%m-%d %H:%M")}').font.size = Pt(9)

# ── SAVE ──
output_path = os.path.expanduser('~/Desktop/BexieMart_Progress_Report_May-26-2026.docx')
doc.save(output_path)
print(f'Report saved to: {output_path}')
print(f'File size: {os.path.getsize(output_path) / 1024:.0f} KB')
